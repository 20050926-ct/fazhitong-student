# -*- coding: utf-8 -*-
"""将 temp/legal-ai-export.md 转为桌面 Word（需 pip install python-docx）。"""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise SystemExit("请先执行: pip install python-docx")

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "temp" / "legal-ai-export.md"
DESKTOP = Path.home() / "Desktop" / "普法AI问答固定话术.docx"
FALLBACK = ROOT / "temp" / "普法AI问答固定问题.docx"


def main() -> None:
    if not MD.is_file():
        raise SystemExit(f"缺少 {MD}，请先运行: npm run export:legal-ai-doc")
    text = MD.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("> "):
            doc.add_paragraph(line[2:].strip())
        else:
            doc.add_paragraph(line)

    DESKTOP.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(DESKTOP))
    except OSError:
        FALLBACK.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(FALLBACK))
        print("[md_to_docx] Desktop save failed; wrote project temp/ (close Desktop docx and re-run)")
        return
    print("[md_to_docx] OK")  # 避免 Windows 控制台编码无法打印中文路径


if __name__ == "__main__":
    main()
